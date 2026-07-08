"""
Génère fiche-modele.docx et memo-modele.docx dans assets/templates/.
Chaque modèle est un gabarit A4 pré-formaté : l'équipe remplace les textes
entre [crochets] par le contenu extrait via Copilot Edge.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

BLEU     = RGBColor(0x1A, 0x3C, 0x5E)
GRIS_TXT = RGBColor(0x55, 0x5C, 0x6E)


# ── helpers ────────────────────────────────────────────────────────────────

def set_a4(doc):
    s = doc.sections[0]
    s.page_width   = Cm(21)
    s.page_height  = Cm(29.7)
    s.left_margin  = Cm(2.5)
    s.right_margin = Cm(2.5)
    s.top_margin   = Cm(2)
    s.bottom_margin = Cm(2)

def r(para, text, size=10, bold=False, italic=False, color=None):
    run = para.add_run(text)
    run.font.name  = "Calibri"
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    return run

def spacing(para, before=0, after=6):
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after  = Pt(after)

def hline(para, color="1A3C5E"):
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "8")
    bot.set(qn("w:space"), "4")
    bot.set(qn("w:color"), color)
    pBdr.append(bot)
    pPr.append(pBdr)

def shade(para, fill="F0F4F8"):
    pPr   = para._p.get_or_add_pPr()
    shd   = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill)
    pPr.append(shd)

def clear(doc):
    for p in doc.paragraphs:
        p._p.getparent().remove(p._element)


# ── fiche animateur ────────────────────────────────────────────────────────

def fiche_animateur(path):
    doc = Document()
    set_a4(doc)
    clear(doc)

    # En-tête
    p = doc.add_paragraph()
    r(p, "[NOM DE L'ATELIER]", size=20, bold=True, color=BLEU)
    spacing(p, after=2)

    p = doc.add_paragraph()
    r(p, "Fiche animateur", size=12, italic=True, color=BLEU)
    hline(p)
    spacing(p, after=10)

    # Déroulé
    p = doc.add_paragraph()
    r(p, "DÉROULÉ", size=12, bold=True, color=BLEU)
    spacing(p, before=2, after=6)

    for i in range(1, 9):
        p = doc.add_paragraph()
        r(p, f"{i} — [Titre de l'étape]", size=11, bold=True)
        spacing(p, before=4, after=2)

        p = doc.add_paragraph()
        r(p, "• [Point clé 1]", size=10)
        p.paragraph_format.left_indent = Cm(0.5)
        spacing(p, after=1)

        p = doc.add_paragraph()
        r(p, "• [Point clé 2]", size=10)
        p.paragraph_format.left_indent = Cm(0.5)
        spacing(p, after=2)

        p = doc.add_paragraph()
        r(p, "⏱ Durée estimée : 2 min", size=9, italic=True, color=GRIS_TXT)
        spacing(p, after=8)

    # Pied de page
    p = doc.add_paragraph()
    hline(p)
    spacing(p, after=4)

    p = doc.add_paragraph()
    r(p, "ARIANE 47 — Conseillers Numériques Département Lot-et-Garonne",
      size=9, bold=True, color=BLEU)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    spacing(p, after=2)

    p = doc.add_paragraph()
    r(p, "05 53 47 31 32 · conseiller-numerique@lotetgaronne.fr",
      size=9, color=GRIS_TXT)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(path)
    print(f"✓ {path}")


# ── mémo participant ───────────────────────────────────────────────────────

def memo_participant(path):
    doc = Document()
    set_a4(doc)
    clear(doc)

    # En-tête
    p = doc.add_paragraph()
    r(p, "[NOM DE L'ATELIER]", size=20, bold=True, color=BLEU)
    spacing(p, after=2)

    p = doc.add_paragraph()
    r(p, "Mémo participant", size=12, italic=True, color=BLEU)
    hline(p)
    spacing(p, after=10)

    # Étapes clés
    p = doc.add_paragraph()
    r(p, "ÉTAPES CLÉS", size=12, bold=True, color=BLEU)
    spacing(p, before=2, after=6)

    for i in range(1, 9):
        p = doc.add_paragraph()
        r(p, f"{i} — [Titre de l'étape]", size=11, bold=True)
        spacing(p, before=4, after=1)

        p = doc.add_paragraph()
        r(p, "→ [Une phrase résumant l'essentiel à retenir]",
          size=10, italic=True, color=GRIS_TXT)
        p.paragraph_format.left_indent = Cm(0.5)
        spacing(p, after=8)

    # Bon à savoir
    p = doc.add_paragraph()
    r(p, "BON À SAVOIR", size=12, bold=True, color=BLEU)
    hline(p)
    spacing(p, before=4, after=6)

    for i in range(1, 4):
        p = doc.add_paragraph()
        shade(p)
        r(p, f"• [Astuce ou information complémentaire {i}]", size=10)
        p.paragraph_format.left_indent = Cm(0.4)
        spacing(p, after=3)

    # Pied de page
    p = doc.add_paragraph()
    spacing(p, before=8, after=4)
    hline(p)

    p = doc.add_paragraph()
    r(p, "05 53 47 31 32 · conseiller-numerique@lotetgaronne.fr",
      size=9, color=GRIS_TXT)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(path)
    print(f"✓ {path}")


# ── main ───────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    os.makedirs("assets/templates", exist_ok=True)
    fiche_animateur("assets/templates/fiche-modele.docx")
    memo_participant("assets/templates/memo-modele.docx")
